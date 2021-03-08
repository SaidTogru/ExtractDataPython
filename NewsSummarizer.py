from selenium import webdriver
import json,time
import pandas
from selenium.webdriver import ActionChains
from selenium.webdriver.common.keys import Keys
from pandas import *
import numpy as np
import urllib.request
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import os
from google_trans_new import google_translator

language=["Deutsch","Chinesisch","Englisch","Französisch", "Italienisch", "Japanisch","Niederländisch","Polnisch","Portugiesisch","Russisch","Spanisch"]

def main(languageNumber):  
    translator = google_translator()  
    language=["de","zh-CN","nl","en","fr", "it", "ja","pl","pt","ru","es","tr","ar"]
    languageName=["German","Chinese","Dutch","English","French", "Italian", "Japanese","Polish","Portugese","Russian","Spansish","Turkish","Arabic"]
    driver=webdriver.Chrome(executable_path=r'C:\chromedriver_win32\chromedriver.exe')
    document = Document()
    document.add_heading("TODAY'S NEWS ("+languageName[languageNumber]+")", 0)
#GOOGLE NEWS
    driver.get("https://news.google.com/topics/CAAqJggKIiBDQkFTRWdvSUwyMHZNRGx1YlY4U0FtVnVHZ0pWVXlnQVAB?hl=en-US&gl=US&ceid=US%3Aen")
    driver.maximize_window() #For maximizing window
    article =driver.find_elements_by_xpath('//*[@id="yDmH0d"]/c-wiz/div/div[2]/c-wiz/div/div[2]/div/main/c-wiz/div/div/main/div[1]/div[*]/div/div/article')
    link =driver.find_elements_by_xpath('//*[@id="yDmH0d"]/c-wiz/div/div[2]/c-wiz/div/div[2]/div/main/c-wiz/div/div/main/div[1]/div[*]/div/div/article/a')
    styles = document.styles
    style = styles.add_style('Überschrift', WD_STYLE_TYPE.PARAGRAPH) 
    style.font.size = Pt(16)
    p = document.add_paragraph('1. GOOGLE NEWS')
    p.style = document.styles['Überschrift']

    for x in range(0, 3):
        p = document.add_paragraph()
        lg=language[languageNumber]
        translation = translator.translate(str((article[x].text).partition('\n')[0]), lang_src="en",lang_tgt=str(language[languageNumber]))
        p.add_run("•   "+translation+" ➜ ").bold = True
        hyperlink = add_hyperlink(p, str(link[x].get_attribute("href")), '[Link]', '0fff13', False)
    driver.quit()
    document.save('News.docx')
 
#AL JAZEERA
    driver=webdriver.Chrome(executable_path=r'C:\chromedriver_win32\chromedriver.exe')
    driver.get("https://www.aljazeera.com/news/")
    driver.maximize_window() #For maximizing window
    article =driver.find_elements_by_xpath('//*[@id="root"]/div/div[4]/div/div[*]/article')
    link =driver.find_elements_by_xpath('//*[@id="root"]/div/div[4]/div/div[*]/article/div[2]/div[1]/h3/a')
    p = document.add_paragraph('2. ALJAZEERA NEWS')
    p.style = document.styles['Überschrift']

    for x in range(0, 3):
        p = document.add_paragraph()
        translation = translator.translate(str((article[x].text).partition('\n')[0]), lang_src="en",lang_tgt=str(language[languageNumber]))
        p.add_run("•   "+translation+" ➜ ").bold = True
        hyperlink = add_hyperlink(p, link[x].get_attribute("href"), '[Link]', 'ffe942', False)
    driver.quit()

#NEW YORK TIMES
    driver=webdriver.Chrome(executable_path=r'C:\chromedriver_win32\chromedriver.exe')
    driver.get("https://www.nytimes.com/section/world")
    driver.maximize_window() #For maximizing window
    article =driver.find_elements_by_xpath('//*[@id="collection-highlights-container"]/div[1]/ol/li[*]/article/div/h2/a')
    link =driver.find_elements_by_xpath('//*[@id="collection-highlights-container"]/div[1]/ol/li[*]/article/div/h2/a')
    p = document.add_paragraph('3. THE NEW YORK TIMES')
    p.style = document.styles['Überschrift']

    for x in range(0, 3):
        p = document.add_paragraph()
        if (x==2):
            articleExc =driver.find_element_by_xpath('//*[@id="collection-highlights-container"]/div[1]/ol/li[3]/ol/li[1]/article/div/h2/a')
            translation = translator.translate(str((articleExc.text).partition('\n')[0]), lang_src="en",lang_tgt=str(language[languageNumber]))
            p.add_run("•   "+translation+" ➜ ").bold = True
            hyperlink = add_hyperlink(p, articleExc.get_attribute("href"), '[Link]', '969696', False)
        else:
            translation = translator.translate(str((article[x].text).partition('\n')[0]), lang_src="en",lang_tgt=str(language[languageNumber]))
            p.add_run("•   "+translation+" ➜ ").bold = True
            hyperlink = add_hyperlink(p, link[x].get_attribute("href"), '[Link]', '969696', False)
    driver.quit()

#BBC
    driver=webdriver.Chrome(executable_path=r'C:\chromedriver_win32\chromedriver.exe')
    driver.get("https://www.bbc.com/news/world")
    driver.maximize_window() #For maximizing window
    article =driver.find_elements_by_xpath('//*[@id="topos-component"]/div[3]/div[2]/div[1]/div/div/div/div[3]/div/div[*]/div/div[2]/div[1]/a')
    link =driver.find_elements_by_xpath('//*[@id="topos-component"]/div[3]/div[2]/div[1]/div/div/div/div[3]/div/div[*]/div/div[2]/div[1]/a')
    p = document.add_paragraph('4. BBC NEWS')
    p.style = document.styles['Überschrift']

    p = document.add_paragraph()
    articleExc=driver.find_element_by_xpath('//*[@id="topos-component"]/div[3]/div[2]/div[1]/div/div/div/div[1]/div/div[2]/div[1]/a')
    p.add_run("•   "+(articleExc.text).partition('\n')[0]+" ➜ ").bold = True   
    hyperlink = add_hyperlink(p, articleExc.get_attribute("href"), '[Link]', 'ff0000', False)

    for x in range(0, 2):
        p = document.add_paragraph()
        translation = translator.translate(str((article[x].text).partition('\n')[0]), lang_src="en",lang_tgt=str(language[languageNumber]))
        p.add_run("•   "+translation+" ➜ ").bold = True    
        hyperlink = add_hyperlink(p, link[x].get_attribute("href"), '[Link]', 'ff0000', False)
    driver.quit()
#TWITTER NEWS
    driver=webdriver.Chrome(executable_path=r'C:\chromedriver_win32\chromedriver.exe')
    driver.get("https://getdaytrends.com/")
    driver.maximize_window() #For maximizing window
    article =driver.find_elements_by_xpath('/html/body/main/div/div[2]/div[2]/section/div/div[1]/table/tbody/tr[*]/td[1]/a')
    p = document.add_paragraph('5. TWITTER NEWS')
    p.style = document.styles['Überschrift']

    for x in range(0, 3):
        p = document.add_paragraph()
        p.add_run("•   "+(article[x].text).partition('\n')[0]+" ➜ ").bold = True    
        hyperlink = add_hyperlink(p, "https://twitter.com/hashtag/"+(article[x].text)[1:]+"?lang=en","[Link]", '1ABEF0', False)
    driver.quit()
    document.save('News.docx')

def add_hyperlink(paragraph, url, text, color, underline):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)
    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)
    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

if __name__ == '__main__':
    main(3)
    #0: GERMAN
    #1: CHINESE
    #2: DUTCH
    #3: ENGLISH
    #4: FRENCH
    #5: ITALIAN
    #6: JAPANESE
    #7: POLISH
    #8: PORTUGESE
    #9: RUSSIAN
    #10: SPANSIH
    #11: TURKISH
    #12: ARABIC
